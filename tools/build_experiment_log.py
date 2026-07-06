"""Generate experiments/EXPERIMENT_LOG.docx — full history of agent versions."""
import os
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

OUT_DIR = os.path.join(os.path.dirname(__file__), "..", "experiments")
os.makedirs(OUT_DIR, exist_ok=True)
OUT_PATH = os.path.join(OUT_DIR, "EXPERIMENT_LOG.docx")
OUT_PATH_TMP = os.path.join(OUT_DIR, "EXPERIMENT_LOG_new.docx")

GREEN  = RGBColor(0x1a, 0x7a, 0x1a)
RED    = RGBColor(0xb0, 0x00, 0x00)
ORANGE = RGBColor(0xb8, 0x5c, 0x00)
BLUE   = RGBColor(0x00, 0x44, 0x99)
BLACK  = RGBColor(0x00, 0x00, 0x00)

BEST_V18 = 0.7977


def set_col_width(table, col_idx, width_inches):
    for row in table.rows:
        row.cells[col_idx].width = Inches(width_inches)


def heading(doc, text, level=1):
    p = doc.add_heading(text, level=level)
    return p


def colored_run(para, text, color, bold=False):
    run = para.add_run(text)
    run.font.color.rgb = color
    run.bold = bold
    return run


def result_color(wr, best=BEST_V18):
    if wr is None:
        return ORANGE
    if wr > best + 0.0005:
        return GREEN
    if wr >= best - 0.0015:
        return ORANGE   # within noise
    return RED


def add_table_row(table, cells, bold_col0=True):
    row = table.add_row()
    for i, (text, color) in enumerate(cells):
        cell = row.cells[i]
        p = cell.paragraphs[0]
        run = p.add_run(text)
        if color:
            run.font.color.rgb = color
        run.bold = (bold_col0 and i == 0)
        run.font.size = Pt(9)
    return row


doc = Document()

# ── Title ──────────────────────────────────────────────────────────────────
title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = title.add_run("PTCG AI Battle — Agent Experiment Log")
r.bold = True
r.font.size = Pt(18)

sub = doc.add_paragraph()
sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
r2 = sub.add_run("Kaggle Competition · June 2026 · Win-rate target: 0.85")
r2.font.size = Pt(11)
r2.font.color.rgb = BLUE

doc.add_paragraph()

# ── Competition Context ────────────────────────────────────────────────────
heading(doc, "Competition Context", 1)
doc.add_paragraph(
    "Simulation comp (pokemon-tcg-ai-battle): submit agent .tar.gz, rated by TrueSkill ladder. "
    "Strategy comp ($240k, 8×$30k): submit ≤2000-word writeup scored 70% model / 20% deck / 10% report. "
    "Engine: cabt (C++ ctypes, cg.dll Windows). CPU-only, Python 3.13.5. "
    "Eval harness: tools/comprehensive_eval.py — 5 seeds × 1000-pool × 5 agents = 25k games (Part 1) "
    "+ mirror 500g × 5 = 2500 (Part 2) + trainer decks 200g × 5 = 1000/deck (Part 3)."
)

# ── Deck (D5, locked) ─────────────────────────────────────────────────────
heading(doc, "Locked Deck: D5 (Ancient Box Fighting)", 1)
doc.add_paragraph(
    "Koraidon ex ×3 (979), Mega Lucario ex ×2 (678), Riolu ×2 (333), "
    "Ogerpon ex ×2 (117), Regirock ex ×2 (447), Fezandipiti ex ×2 (140), "
    "Bloodmoon Ursaluna ×3 (135), Ting-Lu ×2 (41). "
    "Trainers: Ultra Ball ×4, Buddy-Buddy Poffin ×4, Pokégear ×2, Switch ×2, "
    "Night Stretcher ×1, Boss's Orders ×3, Prime Catcher ×1, Lillie's Determination ×2, "
    "Air Balloon ×2, Tarragon ×1, Lively Stadium ×1, Lucky Helmet ×1. "
    "Energy: Fighting ×18."
)

# ── Key Attack Data ────────────────────────────────────────────────────────
heading(doc, "Key Attack Data", 2)
tbl = doc.add_table(rows=1, cols=3)
tbl.style = "Light Grid Accent 1"
hdr = tbl.rows[0].cells
for i, h in enumerate(["Card", "Attack (ID)", "Notes"]):
    p = hdr[i].paragraphs[0]
    r = p.add_run(h)
    r.bold = True
    r.font.size = Pt(9)

attacks = [
    ("Mega Lucario ex (678)", "Aura Jab (982): 130 dmg [F], attach ≤3F from discard to bench\nMega Brave (983): 270 dmg [FF], can't use next turn", "Stage 1 Mega ex"),
    ("Regirock ex (447)",     "Regi Charge (628): 0 dmg [C], attach 2F from discard to self\nGiant Rock (629): 140 + 140 vs Stage 2 [FCCC]", "Energy accelerator"),
    ("Fezandipiti ex (140)",  "Cruel Arrow (183): 100 to ANY opp Pokémon [CCC] (engine reports dmg=0!)", "Darkness type, bench sniper"),
    ("Bloodmoon Ursaluna (135)", "Mad Bite (175): 100 + 30 per damage counter on opp active [FFC]", "Single-prize finisher"),
    ("Koraidon ex (979)",     "Orichalcum Fang (1409): 200 + 120 if our Pokémon KO'd last turn [FFC]", "2-prize finisher"),
]
for card, atk, note in attacks:
    row = tbl.add_row()
    for i, txt in enumerate([card, atk, note]):
        p = row.cells[i].paragraphs[0]
        run = p.add_run(txt)
        run.font.size = Pt(9)

doc.add_paragraph()

# ── Opponents in Eval ─────────────────────────────────────────────────────
heading(doc, "Evaluation Opponents", 2)
doc.add_paragraph(
    "random — legal random actions. "
    "greedy — type-priority attacker (attacks ASAP). "
    "tempo — energy-aware impatient attacker (cheapest attack first). "
    "strong — full heuristic: energy-need aware, danger retreat, lethal detection. "
    "pivot — low-retreat Fighting wall deck; constantly pivots wounded Pokémon to bench."
)

# ── Summary Table ──────────────────────────────────────────────────────────
heading(doc, "All Agent Versions — Summary", 1)

tbl2 = doc.add_table(rows=1, cols=7)
tbl2.style = "Light Grid Accent 1"
tbl2.alignment = WD_TABLE_ALIGNMENT.LEFT
hdrs = ["Version", "Overall", "vs V18", "random", "greedy", "tempo", "strong"]
for i, h in enumerate(hdrs):
    p = tbl2.rows[0].cells[i].paragraphs[0]
    r = p.add_run(h)
    r.bold = True
    r.font.size = Pt(9)

VERSIONS = [
    # (name, overall, random, greedy, tempo, strong, status, note)
    ("V1 (baseline)",  0.7640, None, None, None, None, "kept",     "Initial rules + weight scaffold"),
    ("V2",             0.7640, None, None, None, None, "kept",     "Weight overhaul"),
    ("V3",             0.7692, None, None, None, None, "kept",     "Prize math + urgency scoring"),
    ("V4",             None,   None, None, None, None, "reverted", "Reverted before this session"),
    ("V5",             None,   None, None, None, None, "reverted", "Reverted before this session"),
    ("V6",             0.7721, None, None, None, None, "kept",     "Damage-weighted energy attachment"),
    ("V7",             0.7918, None, None, None, None, "kept",     "Powered switch-in (BREAKTHROUGH)"),
    ("V8",             None,   None, None, None, None, "reverted", "Reverted before this session"),
    ("V9",             None,   None, None, None, None, "reverted", "Reverted before this session"),
    ("V10",            0.7924, None, None, None, None, "kept",     "Evolve priority (+attack-ready bonus)"),
    ("V11",            0.7931, None, None, None, None, "kept",     "Game-winning gust override + bench snipe"),
    ("V12",            0.7945, None, None, None, None, "kept",     "ex-evolution card value fix + EVOLVES_FROM energy awareness"),
    ("V13",            0.7964, None, None, None, None, "kept",     "attach_dmg_scale=0.30, gust_ko=300, recovery item priority"),
    ("V14",            None,   None, None, None, None, "reverted", "Reverted before this session"),
    ("V15",            None,   None, None, None, None, "reverted", "Reverted before this session"),
    ("V16",            0.7921, None, None, None, None, "reverted", "retreat_danger 300→500: Koraidon/Lucario retreated instead of attacking"),
    ("V17",            0.7807, None, None, None, None, "reverted", "Smart retreat + damage bonuses: major regression"),
    ("V18 ★ BEST",     0.7977, 0.8530, 0.8166, 0.7700, 0.7356, "BEST",    "Mad Bite scaling from maxHp-based damage counters"),
    ("V19",            0.7961, None, None, None, None, "reverted", "Orichalcum Fang+Giant Rock+Aura Jab: appearThisTurn too noisy"),
    ("V20",            0.7948, None, None, None, None, "reverted", "Tarragon+Regi Charge scoring: confused agent"),
    ("V21 (buggy)",    None,   None, None, None, None, "reverted", "Boss Orders KO bonus — had NameError bug (state/yi out of scope)"),
    ("V21 (fixed)",    0.7967, 0.8508, 0.8180, 0.7700, 0.7276, "reverted", "Bug fixed, but still -0.0010 vs V18"),
    ("V22",            0.7954, 0.8544, 0.8098, 0.7746, 0.7244, "reverted", "Giant Rock +140 vs Stage 2: -0.0023 vs V18"),
    ("V23a",           0.7958, 0.8526, 0.8134, 0.7684, 0.7300, "reverted", "Cruel Arrow: bench KO scored lethal — premature attacks"),
    ("V24",            0.7908, 0.8516, 0.8062, 0.7696, 0.7160, "reverted", "ATTACH_TO smart Aura Jab targeting — worst regression this session (−0.0069)"),
    ("V25",            0.7958, 0.8522, 0.8080, 0.7746, 0.7256, "reverted", "Orichalcum Fang +120 via prize-count drop detection — still −0.0019 vs V18"),
    # ── Planner phase ──────────────────────────────────────────────────────
    ("PlannerV1",      None,   None, None, None, None, "kept",
     "TurnPlanner shipped: beam=4, depth=6, k=2 determinizations, 0.15s think. "
     "H2H vs rules: 0.650 (z=+2.44, 60 games). Policy routes MAIN decisions to planner."),
    ("PlannerV2 ★",    0.829, 0.862, 0.843, 0.813, 0.787, "BEST",
     "Bench-eval + gust-beam injection. "
     "Added opp_bench_dmg=35 and bench_setup_ko=22 to EVAL. "
     "_candidate_main_options unconditionally includes Boss's Orders / Prime Catcher "
     "when opponent has bench Pokémon. "
     "H2H vs old planner: 0.617 (z=+1.86). H2H vs rules: 0.683 (z=+3.05, 60 games). "
     "Field eval (1000 decks, seed 42, 5000 games, PTCG_ENABLE_PLANNER=1): 0.829 OVERALL "
     "(random=0.862, greedy=0.843, tempo=0.813, strong=0.787, pivot=0.841). "
     "+3.1pp vs rules V18 (0.7977). SHIPPED."),
]

for v in VERSIONS:
    name, overall, rand, greedy, tempo, strong, status, note = v
    wr_str = f"{overall:.4f}" if overall else "—"
    delta_str = ""
    if overall and status not in ("running",):
        d = overall - BEST_V18
        delta_str = f"{d:+.4f}" if overall != BEST_V18 else "baseline"

    color_wr = result_color(overall) if overall else ORANGE
    color_d  = GREEN if (overall and overall > BEST_V18 + 0.0005) else (RED if (overall and overall < BEST_V18 - 0.0015) else ORANGE)
    if status == "BEST":
        color_wr = GREEN
        color_d = GREEN

    rand_str   = f"{rand:.4f}" if rand else "—"
    greedy_str = f"{greedy:.4f}" if greedy else "—"
    tempo_str  = f"{tempo:.4f}" if tempo else "—"
    strong_str = f"{strong:.4f}" if strong else "—"

    row = tbl2.add_row()
    data = [
        (name,       GREEN if "BEST" in status else (RED if status == "reverted" else (ORANGE if status == "running" else BLUE))),
        (wr_str,     color_wr),
        (delta_str,  color_d),
        (rand_str,   BLACK),
        (greedy_str, BLACK),
        (tempo_str,  BLACK),
        (strong_str, BLACK),
    ]
    for i, (txt, col) in enumerate(data):
        cell = row.cells[i]
        p = cell.paragraphs[0]
        run = p.add_run(txt)
        run.font.color.rgb = col
        run.font.size = Pt(9)
        if i == 0 and "BEST" in status:
            run.bold = True

doc.add_paragraph()

# ── Detailed Version Notes ─────────────────────────────────────────────────
heading(doc, "Detailed Version Notes", 1)

DETAILS = [
    ("V7 — Powered Switch-In (BREAKTHROUGH, +0.0278 vs V6)", "kept",
     "Added _powered_switch_value(): when choosing which bench Pokémon to bring active "
     "(after KO or retreat), prefer Pokémon that can already attack (score +120 + dmg×0.3) "
     "over cold ones (partial energy +len×15, low HP −40). "
     "Result: 0.7640→0.7918. Largest single jump of the session. "
     "WHY it works: getting a powered attacker in immediately vs wasting turns with a cold bench Pokémon."),
    ("V11 — Game-Winning Gust Override", "kept",
     "Added 50,000 override when a Boss's Orders gust would give us our last prize. "
     "Also added bench-snipe KO scoring: if Cruel Arrow or similar can KO a benched Pokémon, +200+pv×50. "
     "Result: 0.7924→0.7931."),
    ("V13 — Tuned Weights", "kept",
     "attach_dmg_scale=0.30 (was 0.20), gust_ko=300 (was 250). "
     "Recovery items (Night Stretcher/Tarragon) scored higher when hand ≤4. "
     "Result: 0.7945→0.7964."),
    ("V18 — Mad Bite Damage Counter Estimation (BEST, +0.0013 vs V13)", "BEST",
     "Bloodmoon Ursaluna's Mad Bite does 100 + 30 per damage counter on opp active. "
     "Previously scored as 100 flat (engine damage field). "
     "FIX: estimate counters = (opp maxHp − current hp) / 10. "
     "Added 'opp_counters' to ctx dict in _choose_main (from opp_active.get('maxHp') and opp_hp). "
     "In _score_main: if attack_id == 175 and counters > 0: dmg = 100 + counters × 30. "
     "Enables correct lethal detection when the opp active is already damaged. "
     "Result: 0.7964→0.7977."),
    ("V16 — retreat_danger 300→500 (REGRESSION)", "reverted",
     "Hypothesis: raise retreat danger score so agent retreats doomed Active more. "
     "BUG: retreat = 20 + retreat_danger = 520 > attack_base (380) → Koraidon/Lucario "
     "retreated instead of attacking. Core constraint: retreat score must NEVER beat "
     "a non-lethal attack for our best attackers. retreat_danger 300 means max retreat=320 < 380. "
     "Result: 0.7977→0.7921. Reverted."),
    ("V17 — Smart Retreat + Damage Bonuses (MAJOR REGRESSION)", "reverted",
     "Added bench_best_dmg > my_active_dmg condition + retreat_danger=420. "
     "Still caused wrong retreats — combined with bench attacker detection that "
     "fires for unpowered Pokémon. Result: 0.7977→0.7807. Reverted."),
    ("V19 — Orichalcum Fang + Giant Rock + Aura Jab (REGRESSION)", "reverted",
     "Tried to detect Orichalcum Fang +120 bonus via appearThisTurn flag. "
     "BUG: appearThisTurn fires for BOTH KO AND voluntary retreat — too noisy. "
     "Aura Jab bench energy attachment bonus unclear. Too many changes at once. "
     "Result: 0.7977→0.7961. Reverted."),
    ("V20 — Tarragon + Regi Charge (REGRESSION)", "reverted",
     "Regi Charge scored as attach_active − 20 + min(disc_f,2)×30 — confused agent. "
     "Tarragon +25/F energy in discard hurt tempo. Result: 0.7977→0.7948. Reverted."),
    ("V21 — Boss Orders KO Bonus (BUG + REGRESSION)", "reverted",
     "Hypothesis: Boss's Orders should score higher (+300+pv×50) when opp has KO-able bench target. "
     "CRITICAL BUG: _score_main used 'state' and 'yi' from _choose_main scope → NameError on "
     "every Boss play, caught by try/except, fell back to END. "
     "FIX: passed 'opp_bench' and 'my_active' through ctx dict. "
     "Even after fix: 0.7977→0.7967 (−0.0010, within noise but below threshold). Reverted."),
    ("V22 — Giant Rock Stage 2 Detection (REGRESSION)", "reverted",
     "Added _STAGE2_IDS frozenset (116 Stage 2 IDs) at module level. "
     "In _score_main: elif opt.attack_id == 629 and opp_active_id in _STAGE2_IDS: dmg += 140. "
     "Hypothesis: Giant Rock correctly scores 280 vs Stage 2 targets, enabling lethal detection. "
     "Result: 0.7977→0.7954 (−0.0023). strong matchup dropped 0.7356→0.7244. Reverted. "
     "WHY it might hurt: Stage 2 opponents are rare in the random 1000-deck pool; "
     "the change may have had zero upside while introducing noise in attack selection."),
    ("V23a — Cruel Arrow dmg=0 Fix with Bench KO (REGRESSION)", "reverted",
     "Cruel Arrow (183) has engine dmg=0 but does 100 to ANY opp Pokémon. "
     "V23a fix: check active KO, check bench KO (score lethal_base+pv×50), else dmg=100. "
     "BUG: bench KO path scored 10050 — above EVOLVE (880), causing premature Fezandipiti attacks "
     "instead of setting up Mega Lucario first. Result: 0.7977→0.7958. Reverted."),
    ("V23b — Cruel Arrow dmg=100 fix (REGRESSION)", "reverted",
     "Simplified fix: elif opt.attack_id == 183: dmg = 100. "
     "Lets existing lethal detection handle the active KO via effective_damage. "
     "No bench KO special case — avoids premature attacks. "
     "Non-lethal scores ~440 (vs 380 before). "
     "WHY it still regresses: Fezandipiti ex is rarely the active attacker in D5 — "
     "it's a 2-copy draw/setup piece (CCC cost, Darkness type). The dmg=100 fix raises "
     "its non-lethal attack score from 380 to 440, which is still below ATTACH (500–700), "
     "so setup isn't disrupted — but the rare cases where Fezandipiti is forced active "
     "and the agent now attacks (440) instead of doing something else may be suboptimal. "
     "Result: 0.7977→0.7958. Reverted."),
    ("V24 — ATTACH_TO Smart Aura Jab Targeting (WORST REGRESSION, −0.0069)", "reverted",
     "Hypothesis: when Aura Jab asks which bench Pokémon receive 3 free F energies (ATTACH_TO context), "
     "prioritise needy high-damage attackers (needy=+200 bonus, dmg×0.3 scale). "
     "BUG: _option_pokemon() in the ATTACH_TO branch likely resolves to wrong targets or None — "
     "ATTACH_TO options may reference energy cards or discard slots, not bench Pokémon. "
     "Strong dropped to 0.7160 (−0.0196 vs V18). "
     "Root cause: ATTACH_TO was already handled correctly by the generic _ACQUIRE_CONTEXTS path; "
     "overriding it with a broken _option_pokemon lookup broke energy distribution entirely. "
     "Result: 0.7977→0.7908. Reverted."),
    ("V25 — Orichalcum Fang +120 via Prize-Count Detection (REGRESSION)", "reverted",
     "Koraidon ex attack 1409: 200 + 120 if any of our Pokémon were KO'd during opponent's last turn. "
     "FIX: _opp_prize_cache = [6] at module level; each _choose_main call checks if opp prize count "
     "dropped since last turn → we_were_koed=True → dmg += 120 in _score_main. "
     "WHY it regresses: the prize-count cache is global (module-level), but the eval runs multiple "
     "games sequentially in the same process — cache bleeds between games. Game N's final prize "
     "count contaminates Game N+1's turn-1 detection, causing false positives and premature attacks. "
     "Fix would require resetting cache at game start (no clear hook in the current API). "
     "Result: 0.7977→0.7958. Reverted."),
    # ── Planner phase ──────────────────────────────────────────────────────
    ("PlannerV1 — TurnPlanner (SHIPPED, PROVEN)", "kept",
     "Built TurnPlanner: beam search over MAIN actions using the engine's persistent search tree. "
     "Beam=4, depth=6, k=2 determinizations, 0.15s think time. "
     "Policy.py routes MAIN decisions (select_type==0) to planner; all other decisions fall through to rules. "
     "Eval: prize race + opp_active damage + setup_ko + survivability + board development. "
     "H2H vs rules: 0.650 (z=+2.44, 60 games). E2E crash-free (~50ms/decision)."),
    ("PlannerV2 — Bench Eval + Gust Beam (SHIPPED, BEST)", "BEST",
     "Identified weakness: pivot/wall decks retreat wounded Pokémon before KOs — "
     "planner eval was blind to opponent bench damage, so it never valued spreading damage or gusting. "
     "FIX 1 (eval): Added opp_bench_dmg=35 and bench_setup_ko=22 to EVAL dict. "
     "_eval() now iterates op bench: credits (maxHp-hp)/maxHp * prize_value * 35 per damaged bench Pokémon; "
     "also credits 22 * prize_value when we can KO a bench target with our active. "
     "FIX 2 (beam): _candidate_main_options() now unconditionally injects Boss's Orders (1182) / "
     "Prime Catcher (1088) into the beam when opponent has any bench Pokémon — regardless of TYPE_PRIORITY=3. "
     "This ensures the gust+attack line is always explored. "
     "H2H vs PlannerV1: 0.617 (z=+1.86, 60 games). H2H vs rules: 0.683 (z=+3.05, 60 games). "
     "+3.3pp lift over PlannerV1. SHIPPED as final submission. 7/7 tests pass. "
     "Field eval (1000 decks, seed 42, 5000 games): 0.829 OVERALL (+3.1pp vs rules V18 0.7977). "
     "Per-opponent: random=0.862 (+0.9pp), greedy=0.843 (+2.6pp), tempo=0.813 (+4.3pp), "
     "strong=0.787 (+5.1pp), pivot=0.841."),
]

# V18 is best rules agent (0.7977). PlannerV2 is best overall — 0.683 vs rules H2H (z=+3.05).

for title_text, status, body in DETAILS:
    p = doc.add_paragraph()
    r = p.add_run(title_text)
    r.bold = True
    r.font.size = Pt(11)
    color = GREEN if status == "kept" or status == "BEST" else (RED if status == "reverted" else ORANGE)
    r.font.color.rgb = color

    doc.add_paragraph(body).runs[0].font.size = Pt(10)
    doc.add_paragraph()

# ── Key Constraints & Lessons ──────────────────────────────────────────────
heading(doc, "Key Constraints & Lessons Learned", 1)

lessons = [
    ("retreat_danger constraint",
     "retreat score = w['retreat'] + w['retreat_danger'] = 20 + 300 = 320. "
     "Must NEVER exceed attack_base (380) or the agent retreats instead of attacking. "
     "retreat_danger > 360 → regression. This is dead code for practical purposes."),
    ("appearThisTurn is unusable for Orichalcum Fang",
     "The field fires for BOTH enemy KO AND voluntary retreat. "
     "Cannot distinguish 'our Pokémon was KO'd last turn' from 'we retreated'. Avoid."),
    ("Engine crash cards",
     "Maximum Belt (1158), Hyper Aroma (1082), Grand Tree (1249) all crash cg.dll. Do NOT use."),
    ("Cruel Arrow engine reports dmg=0",
     "attack 183 (Cruel Arrow) has damage=0 in the engine — it's a bench-snipe. "
     "Must manually set dmg=100 in _score_main before lethal detection."),
    ("Boss's Orders / _score_main scope bug pattern",
     "_score_main does not have access to 'state' or 'yi' from _choose_main. "
     "Any board-state data needed in _score_main must be pre-computed and passed via ctx dict."),
    ("Stage 2 detection is narrow",
     "Only ~116 Stage 2 Pokémon in the pool. Giant Rock bonus fires rarely — "
     "not enough signal to move the 25k-game eval. V22 proved this."),
    ("Eval threshold",
     "Keep only if result ≥ V18 (0.7977). Noise floor ≈ ±0.003 (σ≈0.00255 at 25k games)."),
]

tbl3 = doc.add_table(rows=1, cols=2)
tbl3.style = "Light Grid Accent 1"
for i, h in enumerate(["Lesson", "Detail"]):
    p = tbl3.rows[0].cells[i].paragraphs[0]
    r = p.add_run(h); r.bold = True; r.font.size = Pt(9)

for lesson, detail in lessons:
    row = tbl3.add_row()
    for i, txt in enumerate([lesson, detail]):
        p = row.cells[i].paragraphs[0]
        run = p.add_run(txt)
        run.font.size = Pt(9)
        if i == 0:
            run.bold = True

doc.add_paragraph()

# ── Ideas Queue ────────────────────────────────────────────────────────────
heading(doc, "Ideas Queue (not yet tested)", 1)

ideas = [
    ("Aura Jab bench attach priority",
     "After Aura Jab (982, attach ≤3F to bench), ATTACH_TO context uses generic _card_value. "
     "Better: prefer needy Pokémon (needs_energy=True) and higher-damage attackers first. "
     "Would require switching from _card_value to board-state inspection in _choose_cards for ATTACH_TO."),
    ("Regi Charge as setup move",
     "Regi Charge (628): 0 dmg, attaches 2F from discard to Regirock. Scored at 380 (generic attack). "
     "Could be scored higher if Regirock is active and discard has F energy — basically 2 free energy attachments. "
     "Risk: inflating score above ATTACH (700) causes Regirock to stay active as charger instead of being benched."),
    ("Orichalcum Fang via log parsing",
     "Koraidon ex: +120 if our Pokémon was KO'd last turn. Can't use appearThisTurn (too noisy). "
     "Possible: read obs_dict['logs'] for LogType.KO events from last turn affecting our player. "
     "Would require parsing the log list in _choose_main before building ctx."),
    ("Lucky Helmet awareness",
     "Lucky Helmet (1156): when Active is hit, draw 2 cards. "
     "Currently just plays as 'play_other' (540). Could prioritize attaching when Active is in danger "
     "and hand is small — but the tool/attach timing makes this complex."),
    ("Night Stretcher / Tarragon timing",
     "Recovery items should be prioritized after mid-game (discard pile likely has energy/Pokémon). "
     "Current: +20 when hand ≤4. Could tie to discard pile size (deckCount drop proxy)."),
]

for i, (title_text, detail) in enumerate(ideas):
    p = doc.add_paragraph(style="List Number")
    r = p.add_run(title_text + ": ")
    r.bold = True
    r.font.size = Pt(10)
    p.add_run(detail).font.size = Pt(10)

doc.add_paragraph()

# ── Save ───────────────────────────────────────────────────────────────────
doc.save(OUT_PATH_TMP)
import shutil, os as _os
try:
    _os.replace(OUT_PATH_TMP, OUT_PATH)
    print(f"Saved → {OUT_PATH}")
except PermissionError:
    print(f"Saved → {OUT_PATH_TMP}  (close the old DOCX in Word to auto-replace)")
